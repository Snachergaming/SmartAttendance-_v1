import os
import markdown
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Configuration
CONTENT_DIR = "report_gen/content"
SOURCE_MD_PATH = "project-report/ATTENDRO_PROJECT_REPORT.md"
DIAGRAMS_MARKDOWN_DIR = "docs/diagrams"
OUTPUT_DOCX = "project-report/FINAL_OUTPUT/Attendro_Final_Report.docx"

# Mapping of Diagrams (From Figure List to File)
# 1. System Architecture Diagram – Figure 1 -> docs/diagrams/01-system-architecture.md
# 2. Database Schema (ER Diagram) – Figure 2 -> docs/diagrams/02-database-schema.md
# 3. User Workflow & Experience – Figure 3 -> docs/diagrams/03-user-workflow.md
# 4. Device Interface & Wiring – Figure 4 -> docs/diagrams/04-device-interface.md
# 5. Security Model & Session Token Flow – Figure 5 -> docs/diagrams/05-security-model.md
# 6. Use Case Diagram – Figure 6 -> WE NEED TO CREATE/FIND THIS (Was listed in HTML folder as 06-usecase-diagram.html)
# 7. Context-Level DFD – Figure 7 -> WE NEED TO FIND THIS

# We will extract ASCII diagrams from the MD files in docs/diagrams and insert them into the Word doc as code blocks
# Since we can't easily generate images from ASCII in python without heavy deps, we will format them as Monospace text.

DIAGRAM_MAP = {
    "Figure 1": "docs/diagrams/01-system-architecture.md",
    "Figure 2": "docs/diagrams/02-database-schema.md",
    "Figure 3": "docs/diagrams/03-user-workflow.md",
    "Figure 4": "docs/diagrams/04-device-interface.md",
    "Figure 5": "docs/diagrams/05-security-model.md",
    # Mappings for 6 & 7 (If found in text, we might skip or grab generic)
}

def setup_document_styles(doc):
    # Setup Normal Style (TNR, 12pt, Double Space, Justified)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    p_format = style.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p_format.space_after = Pt(12)
    
    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.25)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(1.25)
        
    # Styles for Headings
    # Heading 1: TNR 14 Capital Bold Center
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman') 
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.page_break_before = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2: TNR 12 Capital Bold Left
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.page_break_before = False
    
    # Heading 3: TNR 12 Bold Normal Left
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(12)

    # Monospace for Diagrams
    styles = doc.styles
    try:
        code_style = styles.add_style('CodeBlock', 1) # 1 is Paragraph style type
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9) # Smaller for diagrams to fit
        code_style.paragraph_format.line_spacing = 1.0 # Single space for diagrams
        code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except:
        pass # Style might exist

def extract_diagram_text(filepath):
    """
    Reads a markdown file and extracts the first code block content (which usually contains the ASCII diagram).
    """
    if not os.path.exists(filepath):
        return None
        
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Regex to find content inside ``` ... ```
    match = re.search(r'```(.*?)```', content, re.DOTALL)
    if match:
        return match.group(1).strip()
    return None

def process_markdown_content(md_content):
    # Same as before, promote headers
    lines = md_content.split('\n')
    processed_lines = []
    
    for line in lines:
        if re.match(r'^###\s+\d+\.\d+\s+', line):
            line = line.replace('###', '##', 1)
        
        # Determine if we need to insert a diagram here?
        # The user report text doesn't explicitly have markers like [Insert Figure 1] usually.
        # But we can look for the "List of Figures" section and maybe insert them in relevant chapters?
        # OR we can append a "Diagrams" section at the end if the user didn't specify location.
        # usually diagrams go into chapters. e.g. Arch in Ch 5.
        
        # Let's simple check if the line mentions "Figure X" and if we have a match, insert it after the paragraph?
        # But text might say "See Figure 1".
        
        processed_lines.append(line)
        
    return processed_lines

def parse_structure(lines):
    structure = []
    current_section = {"title": "Start", "type": "text", "content": []}
    
    # We will split strictly by Headers
    for line in lines:
        # Title Page Marker
        if "## Title Page" in line:
            structure.append(current_section)
            current_section = {"title": "Title Page", "type": "special", "content": []}
        elif "## Certificate" in line:
            structure.append(current_section)
            current_section = {"title": "Certificate", "type": "special", "content": []}
        elif "## Acknowledgement" in line:
            structure.append(current_section)
            current_section = {"title": "Acknowledgement", "type": "special", "content": []}
        elif "## Table of Contents" in line or "## Index" in line:
             structure.append(current_section)
             current_section = {"title": "Table of Contents", "type": "special", "content": []}
        elif "## Abstract" in line:
             structure.append(current_section)
             current_section = {"title": "Abstract", "type": "special", "content": []}
        elif "## List of Figures" in line:
             structure.append(current_section)
             current_section = {"title": "List of Figures", "type": "special", "content": []}
        elif "## List of Tables" in line:
             structure.append(current_section)
             current_section = {"title": "List of Tables", "type": "special", "content": []}
        elif line.startswith("# Chapter"):
            structure.append(current_section)
            title = line.strip("# ").strip()
            current_section = {"title": title, "type": "chapter", "content": []}
        elif line.startswith("# References"):
            structure.append(current_section)
            current_section = {"title": "References", "type": "ref", "content": []}
        else:
            current_section['content'].append(line)
            
    structure.append(current_section)
    return structure

def add_title_page(doc):
    # Manual creation of Title Page Elements
    doc.add_paragraph("") # Spacer
    doc.add_paragraph("") # Spacer
    
    p = doc.add_paragraph("ATTENDRO: Smart Biometric + App-Based Attendance Management System using AI & IoT")
    p.style = doc.styles['Heading 1'] # Centered Bold 14
    # Manually override size to 16 if needed, but 14 is compliant.
    
    doc.add_paragraph("\n\n")
    
    p = doc.add_paragraph("A Project Report Submitted by")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[STUDENT NAME]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\nIn partial fulfillment of the requirements for the Diploma in").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("APPLIED AI & ML")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("\nAt").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Rajarambapu Institute of Technology, Islampur")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("2025–2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n\nUnder the Guidance of").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[GUIDE NAME]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def add_certificate(doc):
    h = doc.add_paragraph("CERTIFICATE")
    h.style = doc.styles['Heading 1']
    
    doc.add_paragraph("\n")
    p = doc.add_paragraph("This is to certify that the project titled “ATTENDRO: Smart Biometric + App-Based Attendance Management System using AI & IoT” has been carried out by [Student Name] under my guidance and supervision in partial fulfillment of the requirements for the award of the Diploma in Applied AI & ML at Rajarambapu Institute of Technology, Islampur, during the academic year 2025–2026.")
    
    doc.add_paragraph("\n\n\n")
    
    # Signatures
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    # Format signatures
    row = table.rows[0]
    row.cells[0].text = "___________________\nGuide"
    row.cells[1].text = "___________________\nH.O.D."
    row.cells[2].text = "___________________\nPrincipal"
    
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    doc.add_page_break()

def generate_docx():
    doc = Document()
    setup_document_styles(doc)
    
    with open(SOURCE_MD_PATH, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    lines = process_markdown_content(md_content)
    structure = parse_structure(lines)
    
    for section in structure:
        title = section['title']
        if title == "Start": continue
        
        if title == "Title Page":
            add_title_page(doc)
            continue
        elif title == "Certificate":
            add_certificate(doc)
            continue
        
        # Normal Section Handling
        # Title as Heading 1
        h1 = doc.add_paragraph(title.upper(), style='Heading 1')
        
        # Content Parsing
        content_lines = section['content']
        
        # Insert Logic: Check for Figure references in text or headers
        # We process line by line to support headers
        for line in content_lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith("## "): 
                # Promoted H2
                h2_text = line.replace("##", "").strip().upper()
                doc.add_paragraph(h2_text, style='Heading 2')
                
                # Check if we should insert a diagram under specific headings?
                # E.g. Under "4.1 System Overview" (Chapter 4) -> Insert Figure 1
                # E.g. Under "5.2 Database..." -> Insert Figure 2
                
                if "SYSTEM OVERVIEW" in h2_text and "4" in title: # Ch 4, Sysem Overview
                    dia_text = extract_diagram_text("docs/diagrams/01-system-architecture.md")
                    if dia_text:
                        doc.add_paragraph("Figure 1: System Architecture", style='Caption')
                        p = doc.add_paragraph(dia_text, style='CodeBlock')
                
                if "DATABASE" in h2_text and "5" in title:
                     dia_text = extract_diagram_text("docs/diagrams/02-database-schema.md")
                     if dia_text:
                        doc.add_paragraph("Figure 2: Database Schema", style='Caption')
                        p = doc.add_paragraph(dia_text, style='CodeBlock')

            elif line.startswith("### "):
                # H3
                h3_text = line.replace("###", "").strip()
                doc.add_paragraph(h3_text, style='Heading 3')
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line.lstrip("- *"), style='List Bullet')
            elif bool(re.match(r'^\d+\.', line)):
                 p = doc.add_paragraph(line, style='List Number') # This often needs cleaner parsing for just the text
            else:
                 # Standard Paragraph
                 # Check for "Figure" mentions to insert diagrams inline
                 # But ASCII diagrams start new blocks.
                 p = doc.add_paragraph(line, style='Normal')
                 
                 # Figure 3: User Workflow - Usually in Ch 5 or Ch 4
                 if "User Workflow" in title or "Methodology" in title:
                     if "Figure 3" in line and not "List of Figures" in title:
                         dia_text = extract_diagram_text("docs/diagrams/03-user-workflow.md")
                         if dia_text:
                             doc.add_paragraph("Figure 3: User Workflow", style='Caption')
                             doc.add_paragraph(dia_text, style='CodeBlock')

                 # Figure 4: Device Interface - Ch 5 Hardware Design
                 if "Hardware Design" in line:
                      dia_text = extract_diagram_text("docs/diagrams/04-device-interface.md")
                      if dia_text:
                          doc.add_paragraph("Figure 4: Device Interface", style='Caption')
                          doc.add_paragraph(dia_text, style='CodeBlock')
        
    
    doc.save(OUTPUT_DOCX)
    print(f"DOCX Generated: {OUTPUT_DOCX}")

if __name__ == "__main__":
    generate_docx()
